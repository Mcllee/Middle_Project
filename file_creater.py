from Window_set import *    # 기본 윈도우 파일
from tkinter import filedialog

import Main_Page as mp      # 메인 페이지 파일
import Log_in_Page as lip   # 로그인 페이지 파일

import docx
from docx import Document   # 워드 저장
import datetime             # 텍스트 현재시간 기록
import os                   # 파일 디렉토리 접근


memo_txt = Text(win_root, width=115, height=40)
menu = Menu()
menu_file = Menu(menu, tearoff=False)


def delete_fc(next_page):
    memo_txt.place_forget()
    menu_file.delete(0, END)
    win_root.config(menu=False)

    match (next_page):  # 이동 횟수(깊이)에 따라 넘버링
        case 1:
            lip.set_log_in_page("")
        case 2:
            mp.Main_Page_set()


def save_file():
    file_name = filedialog.asksaveasfile(parent=win_root, title="파일 선택기",
                                         defaultextension=".*",
                                         filetypes=(("텍스트 파일(.txt)", "*.txt"),
                                                    ("워드 파일(.docx)", "*.docx"),
                                                    ("모든 파일", "*.*")
                                                    )
                                        )
    # 다이얼로그 강제종료면 None 반환
    if file_name == None:
        return

    # 파일의 생성 시간을 받는다.
    file_time = datetime.datetime.fromtimestamp(os.path.getctime(file_name.name)).strftime('\n\n%Y년 %m월 %d일 %H시 %M분에 기록 되었습니다.')

    # 파일 형식에 따라서 저장 방법이 다르다.
    match file_name.name.split('.')[1]:
        case 'docx':
            doc = Document()
            doc.add_paragraph(memo_txt.get("1.0", END))
            doc.add_paragraph(file_time)
            doc.save(file_name.name)
        case 'txt':
            save_text = open(file_name.name, 'w')
            save_text.write(memo_txt.get("1.0", END))
            save_text.write(file_time)
            save_text.close()


def open_file():
    file_name = filedialog.askopenfilename(title='select a text file',
                                           filetypes=(("텍스트 파일(.txt)", "*.txt"),
                                                      ("워드 파일(.docx)", "*.docx"),
                                                      ("모든 파일", "*.*")))
    # 다이얼로그 강제종료면 공백 반환
    if file_name == '':
        return

    match file_name.split('.')[1]:
        case 'txt':
            save_file = ''
            try:
                save_file = open(file_name, 'r')
                memo_txt.delete("1.0", END)  # 출력부 미리 비우기
                memo_txt.insert(END, save_file.read())
            except UnicodeDecodeError:  # 인코딩 방식을 변경한다.
                save_file.close()
                save_file = open(file_name, 'r', encoding='UTF-8')
                memo_txt.delete("1.0", END)  # 출력부 미리 비우기
                memo_txt.insert(END, save_file.read())
            except FileNotFoundError:
                pass
        case 'docx':
            save_file = docx.Document(file_name)

            memo_txt.delete("1.0", END)  # 출력부 미리 비우기
            for _, paragraph in enumerate(save_file.paragraphs):
                memo_txt.insert(END, paragraph.text)


class file_creater:
    #  윈도우 위젯 선언

    @staticmethod
    def set_fc(input_name):
        mp.page_address.delete(0, END)
        mp.page_address.insert(0, f"파일 제작소에 어서오세요! {input_name}님!")

        mp.back_page_btn.config(command=lambda: [delete_fc(1)])
        mp.forward_page_btn.config(command=DISABLED)

        memo_txt.place(x=100, y=100)

        menu_file.add_command(label='불러오기', command=open_file, accelerator='Ctrl+O')
        menu_file.add_command(label='저장하기', command=save_file, accelerator='Ctrl+S')
        menu_file.add_separator()
        menu_file.add_command(label='종료', command=lambda: [delete_fc(1)], accelerator='Ctrl+Q')
        menu.add_cascade(label='File', menu=menu_file)

        win_root.config(menu=menu)
